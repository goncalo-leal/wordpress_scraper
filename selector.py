import re
import sys
import json
import getopt
import shutil, os
from docx import Document
from datetime import datetime

months = {
    "Janeiro": ["January", 31],
    "Fevereiro": ["February", 28],
    "Março": ["March", 31],
    "Abril": ["April", 30],
    "Maio": ["May", 31],
    "Junho": ["June", 30],
    "Julho": ["July", 31],
    "Agosto": ["August", 31],
    "Setembro": ["September", 30],
    "Outubro": ["October", 31],
    "Novembro": ["November", 30],
    "Dezembro": ["December", 31]
}

def print_help():
    print('Limit by number: selector.py -n|--n <number of poems> -d|--dir <output directory>')
    print('Total of poems: selector.py -t|--total')
    print('Date of poem n: selector.py -x|--x <index of poem>')
    print('Limit by date: selector.py -l|--date')

def get_poems():
    f = open('./main.json')
    poems = json.load(f)
    f.close()

    return poems

def get_total():
    poems = get_poems()

    print(f"Total: {len(poems)}")

def order_by_date(poems):
    print("Ordering by date...")

    new_dict = {}
    for poem in poems:
        date_time_str = poems[poem]["date_division"]
        date = date_time_str.split(" ")
        date_time_obj = datetime.strptime(months[date[0]][0] + " " + date[1] + " " + date[2], '%B %d, %Y').timestamp()

        new_dict[poem] = [int(round(date_time_obj)), poems[poem]["poem_file"]]

    sorted_dict = {k: v for k, v in sorted(new_dict.items(), key=lambda item: item[1][0])}

    return sorted_dict

def copy_to_destiny(poems, destiny):
    print("Copying files to destiny directory...")

    for poem in poems:
        shutil.copy(f"./poems/{poems[poem]['file']}", destiny)

def create_word_file(poems, destiny):
    print("Creating word file...")

    document = Document()
    for poem in poems:            
        with open(f"./poems/{poems[poem]['file']}", "r") as read_file:
            document.add_paragraph(read_file.read().strip() + "\n\n")

        date = poems[poem]["date"].split(" ")
        day = date[1].replace(',', '')
        month = list(months.keys()).index(date[0])
        month = f"0{month}" if month < 10 else month

        document.add_paragraph(f"{day}-{month}-{date[2]}")
        document.add_page_break()

    document.save(f'{destiny}.docx')

def limit_date():
    print("Limit by date")

    poems = get_poems()
    sorted_poems = order_by_date(poems)

    print("Enter the date in the format mm/yyyy (ex: 01/2022)")
    date = input()

    while (re.search("[0|1][0-9]/[1|2][0|9][0-9][0-9]", date) == None and date != ""):
        print("Invalid date! Format: mm/yyyy (ex: 01/2022)")
        date = input()

    if date == "":
        print("No date entered")
        sys.exit(2)

    date = date.split("/")
    month = list(months.keys())[int(date[0])-1]
    date_time_obj = datetime.strptime(f"{months[month][0]} {str(months[month][1])} {date[1]}", '%B %d %Y').timestamp()

    to_save = {}
    for poem in sorted_poems.keys():
        if (sorted_poems[poem][0] > date_time_obj):
            break

        to_save[poem] = {
            'name': poem,
            'date': poems[poem]["date_division"],
            'file': poems[poem]["poem_file"]
        }

    create_word_file(to_save, f"{month}_{date[1]}")
    print(f".docx file created with the name: {month}_{date[1]}.docx")

def main(argv):
    output_dir = ''
    n = 0
    try:
        opts, args = getopt.getopt(argv,"lhtx:n:d:",["n=","dir=", "help", "total", "x=", "date"])
    except getopt.GetoptError:
        print_help()
        sys.exit(2)

    if len(opts) == 0:
        print_help()
        sys.exit(2)

    for opt, arg in opts:
        if opt == '-h':
            print_help()
            sys.exit()

        elif opt in ("-n", "--n"):
            if not arg.isdigit():
                print("Number of poems must be a number")
                sys.exit(2)
            n = int(arg)

        elif opt in ("-d", "--dir"):
            output_dir = arg
            
        elif opt in ("-t", "--total"):
            get_total()
            sys.exit(0)

        elif opt in ("-x", "--x"):
            if not arg.isdigit():
                print("Index of poem must be a number")
                sys.exit(2)

            poems = get_poems()
            sorted_poems = order_by_date(poems)
            keys = list(sorted_poems.keys())
            print(poems[keys[int(arg)]]["date_division"])
            sys.exit(0)

        elif opt in ("-l", "--date"):
            limit_date()
            sys.exit(0)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    poems = get_poems()

    sorted_dict = order_by_date(poems)

    to_save = {}
    i = 1
    for poem in sorted_dict.keys():
        to_save[poem] = {
            'name': poem,
            'date': poems[poem]["date_division"],
            'file': poems[poem]["poem_file"]
        }
        
        if i == n:
            print(i)
            break

        i += 1

    copy_to_destiny(to_save, output_dir)

    with open("sorted.json", "w") as write_file:
        json.dump(to_save, write_file, indent=4, ensure_ascii=False)

    create_word_file(to_save, output_dir)

if __name__ == "__main__":
   main(sys.argv[1:])